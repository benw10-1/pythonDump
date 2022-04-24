from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QApplication,QWidget, QVBoxLayout, QPushButton, QFileDialog , QLabel, QTextEdit
import sys
import string, random
from PyQt5.QtGui import QPixmap
from docx import Document
from docx.shared import Inches


def clear_list(l):
    out = []
    for ll in l:
        if ll is None:
            continue
        if isinstance(ll, list):
             ll = clear_list(ll)
             if not ll:
                 continue
        out.append(ll)
    return out



picts = [None]*500
txt = [None]*500
used = [False]*500
filenames = [None]*500


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 603)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.comboBox = QtWidgets.QComboBox(self.centralwidget)
        self.comboBox.setGeometry(QtCore.QRect(30, 20, 201, 22))
        self.comboBox.setObjectName("comboBox")
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setGeometry(QtCore.QRect(0, 390, 591, 161))
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 589, 159))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents)
        self.verticalLayout.setObjectName("verticalLayout")
        self.Inpuptext = QtWidgets.QTextEdit(self.scrollAreaWidgetContents)
        self.Inpuptext.setObjectName("Inpuptext")
        self.verticalLayout.addWidget(self.Inpuptext)
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.pixmaper = QtWidgets.QLabel(self.centralwidget)
        self.pixmaper.setGeometry(QtCore.QRect(610, 400, 171, 141))
        self.pixmaper.setObjectName("pixmaper")
        self.imgchooser = QtWidgets.QPushButton(self.centralwidget)
        self.imgchooser.setGeometry(QtCore.QRect(650, 360, 91, 23))
        self.imgchooser.setObjectName("imgchooser")
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setGeometry(QtCore.QRect(490, 100, 201, 61))
        self.pushButton.setObjectName("pushButton")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(550, 65, 81, 21))
        self.label.setObjectName("label")

        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setGeometry(QtCore.QRect(490, 220, 113, 20))
        self.lineEdit.setObjectName("lineEdit")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(490, 190, 151, 16))
        self.label_2.setObjectName("label_2")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.comboBox.addItem("Question 1")

        self.imgchooser.clicked.connect(self.pic)
        self.comboBox.activated.connect(self.holder)
        self.Inpuptext.textChanged.connect(self.texter)
        self.pushButton.clicked.connect(self.mainrandom)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Randomizer"))
        self.imgchooser.setText(_translate("MainWindow", "Upload Image"))
        self.pushButton.setText(_translate("MainWindow", "Start"))
        self.label.setText(_translate("MainWindow", "Start Randomizer"))
        self.label_2.setText(_translate("MainWindow", "Set output name"))


    def mainrandom(self):
        y = 0
        newl1 = clear_list(txt)
        newl2 = clear_list(filenames)
        randomtup = []
        print (y)
        for r in newl1:
            try:
                newl2[y]
            except:
                newl2.append("")
            try:
                newl1[y]
            except:
                newl1.append("")
            randomtup.append((newl1[y], newl2[y]))
            y+=1
            print (txt[y])
        random.shuffle(randomtup)
        doc = Document()
        i = 0

        doc.add_heading("Test")

        for p in randomtup:
            x = doc.add_paragraph("")
            x.add_run(str(i+1)+". ").bold = True
            x.add_run(p[0])
            if p[1]:
                doc.add_picture(p[1], width=Inches(.5), height=Inches(.5), new=x)
            i+=1
        doc.save("demo.docx")


    def pic(self):
        filename, _ = QtWidgets.QFileDialog.getOpenFileName(None, "Select Image", "", "Images (*.jpg .webm .png)")
        print(filename)
        pixmap = QPixmap(filename)
        pixmap2 = pixmap.scaled(self.pixmaper.width(),self.pixmaper.height())

        self.pixmaper.setPixmap(pixmap2)

        picts[self.comboBox.currentIndex()] = pixmap2
        filenames[self.comboBox.currentIndex()] = filename



    def holder(self):

        print("yee")
        if picts[self.comboBox.currentIndex()] != None:
            self.pixmaper.setPixmap(picts[self.comboBox.currentIndex()])

        if picts[self.comboBox.currentIndex()] == None:
            self.pixmaper.clear()

        if txt[self.comboBox.currentIndex()] != None or "":
            self.Inpuptext.setText(txt[self.comboBox.currentIndex()])

        if txt[self.comboBox.currentIndex()] == None or "":
            self.Inpuptext.setText("")

        print(txt[self.comboBox.currentIndex()])



    def texter(self):
        if not used[self.comboBox.currentIndex()]:
            if self.Inpuptext.toPlainText() != "" or None:
                if self.comboBox.currentIndex() == 0:
                    used[self.comboBox.currentIndex()] = True
                    self.comboBox.addItem("Question 2")

                else:

                    used[self.comboBox.currentIndex()] = True
                    b = self.comboBox.currentIndex() + 2
                    self.comboBox.addItem(f"Question {b}")
        txt[self.comboBox.currentIndex()] = self.Inpuptext.toPlainText()
        print (str(self.comboBox.currentIndex())+ ". "+txt[self.comboBox.currentIndex()])

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
