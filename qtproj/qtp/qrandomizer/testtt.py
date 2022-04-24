from docx import Document
from docx.shared import Inches

import string,random

document = Document("SAMPLE QUIZ FORMAT.docx")
doc1 = Document()
ar1 = []
ar2 = []
num = 2
for x in document.paragraphs:
    ar1.append(x)
    for y in x.runs:
        if y.bold:
            ar2.append(ar1)
            ar1.clear()

for x in range(num):
    random.shuffle(ar2)
    print(doc1.paragraphs(ar2))
    
doc1.save("demo2.docx")




#records = (
    #(3, '101', 'Spam'),
    #(7, '422', 'Eggs'),
    #(4, '631', 'Spam, spam, eggs, and spam')
#)

#table = document.add_table(rows=1, cols=3)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#for qty, id, desc in records:
    #row_cells = table.add_row().cells
    #row_cells[0].text = str(qty)
    #row_cells[1].text = id
    #row_cells[2].text = desc



